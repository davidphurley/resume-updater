def process_work_experience(text):
    """
    Process the WORK EXPERIENCE section so that each '-' line is a bullet point.
    Returns a list of (is_bullet, text) tuples for updating the docx.
    """
    lines = [l.rstrip() for l in text.split('\n')]
    processed = []
    for line in lines:
        if line.strip().startswith('-'):
            processed.append((True, line.lstrip('-').strip()))
        else:
            processed.append((False, line.strip()))
    return processed

def update_work_experience(doc, indices, processed_lines):
    """
    Update the work experience section in the docx.
    Bullet points for lines marked as bullets, regular text otherwise.
    """
    for i, idx in enumerate(indices):
        para = doc.paragraphs[idx]
        para.clear()
        if i < len(processed_lines):
            is_bullet, text = processed_lines[i]
            run = para.add_run(text)
            if is_bullet:
                para.style = 'List Bullet'
            else:
                para.style = None
        else:
            para.text = ''
def process_core_competencies(text):
    """
    Process the CORE COMPETENCIES section to ensure each bullet is 1-2 words and bolded.
    Returns a list of (text, formatting) tuples for updating the docx.
    """
    import re
    # Split into lines, filter out empty, and keep only 1-2 word entries
    lines = [l.strip('-•* 	') for l in text.split('\n') if l.strip()]
    bullets = []
    for line in lines:
        # Only keep 1-2 word entries
        words = line.split()
        if 1 <= len(words) <= 2:
            bullets.append(' '.join(words))
    return bullets

def update_core_competencies(doc, indices, bullets):
    """
    Update the skills section in the docx with bolded bullet points (1-2 words each).
    Each paragraph in indices is replaced with a bullet, bolded.
    """
    from docx.oxml.ns import qn
    for i, idx in enumerate(indices):
        para = doc.paragraphs[idx]
        para.clear()
        if i < len(bullets):
            run = para.add_run(f"• {bullets[i]}")
            run.bold = True
            # Optionally, set style to match original (if needed)
        else:
            para.text = ''
def extract_section_text(doc, indices, word_limit=None):
    """
    Extract text from a list of paragraph indices. Optionally limit total word count.
    Returns a single string with paragraphs joined by newlines.
    """
    paras = [doc.paragraphs[i].text for i in indices]
    if word_limit is not None:
        words = []
        for para in paras:
            for w in para.split():
                if len(words) < word_limit:
                    words.append(w)
        # Rebuild paragraphs with limited words
        out = []
        idx = 0
        for para in paras:
            para_words = para.split()
            if idx + len(para_words) > word_limit:
                out.append(' '.join(words[idx:word_limit]))
                break
            else:
                out.append(' '.join(words[idx:idx+len(para_words)]))
                idx += len(para_words)
        return '\n'.join(out)
    return '\n'.join(paras)

def update_section_text(doc, indices, new_text):
    """
    Update the text of paragraphs at the given indices with new_text (split by lines).
    Preserves formatting by updating only the text in each run.
    """
    lines = new_text.split('\n')
    for i, idx in enumerate(indices):
        if i < len(lines):
            para = doc.paragraphs[idx]
            new_line = lines[i]
            if para.runs:
                start = 0
                for run in para.runs:
                    run_len = len(run.text)
                    run.text = new_line[start:start+run_len]
                    start += run_len
                if start < len(new_line) and para.runs:
                    para.runs[-1].text += new_line[start:]
            else:
                para.text = new_line
        # If fewer lines, leave extra paragraphs unchanged

def parse_docx_sections(doc):
    """
    Parse the docx into logical sections: intro, skills, work experience, and others.
    Returns a dict with section names as keys and lists of paragraph indices as values.
    """
    section_map = {
        'intro': [],
        'skills': [],
        'work_experience': [],
        'other': []
    }
    # Section header keywords (case-insensitive)
    intro_header = 'DATA SCIENTIST'
    skills_header = 'CORE COMPETENCIES'
    work_header = 'WORK EXPERIENCE'
    # Track section starts
    intro_start = skills_start = work_start = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        if intro_header in text:
            intro_start = i
        elif text == skills_header:
            skills_start = i
        elif text == work_header:
            work_start = i
    # Find section ranges
    n = len(doc.paragraphs)
    # Intro: first paragraph after intro_header (just one paragraph)
    if intro_start is not None:
        # Only the next paragraph is the intro content
        if intro_start + 1 < n:
            section_map['intro'] = [intro_start + 1]
    # Skills: first paragraph after skills_header until work_header
    if skills_start is not None and work_start is not None:
        section_map['skills'] = list(range(skills_start+1, work_start))
    # Work experience: from first paragraph after work_header until next all-caps header or end
    if work_start is not None:
        # Find next all-caps header after work_start
        next_header = n
        for i in range(work_start+1, n):
            t = doc.paragraphs[i].text.strip()
            if t.isupper() and len(t) > 3:
                next_header = i
                break
        section_map['work_experience'] = list(range(work_start+1, next_header))
    # All other paragraphs
    used = set(section_map['intro'] or []) | set(section_map['skills'] or []) | set(section_map['work_experience'] or [])
    section_map['other'] = [i for i in range(n) if i not in used]
    return section_map
import os  # For file path operations
import sys  # For command-line arguments
import getpass  # For securely entering API keys
import docx  # For working with docx files
from docx import Document  # Main class for docx manipulation
from docx.shared import Pt  # For font size (not used directly here)
import requests  # For making HTTP requests to Gemini API
from tempfile import NamedTemporaryFile  # For temporary file handling
from docx2pdf import convert  # For converting docx to PDF

def prompt_file_path(prompt_text):
    """
    Prompt the user for a file path until a valid file is provided.
    """
    # Deprecated: now using command-line arguments
    raise NotImplementedError("Use command-line arguments for file paths.")

def read_docx_text(docx_path):
    """
    Read all text from a DOCX file and return both the text and the Document object.
    """
    doc = Document(docx_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text, doc

def read_txt_file(txt_path):
    """
    Read and return the contents of a text file.
    """
    with open(txt_path, 'r', encoding='utf-8') as f:
        return f.read()

def get_gemini_api_key():
    """
    Prompt the user to enter their Gemini API key securely.
    """
    print("You need a Gemini API key. Get one at https://aistudio.google.com/app/apikey")
    return getpass.getpass("Enter your Gemini API key (input hidden): ")

def call_gemini(api_key, resume_text, job_desc):
    """
    Call the Gemini API to rewrite the resume based on the job description.
    Returns the improved resume text.
    """
    url = "https://generativelanguage.googleapis.com/v1/models/gemini-1.5-pro:generateContent?key=" + api_key
    # Construct the prompt for the AI model
    prompt = f"""
You are an expert resume writer. Given the following resume and job description, rewrite the resume to best match the job description, keeping the original formatting and structure as much as possible. Only update the content where relevant. Return the improved resume content only.

Resume:
{resume_text}

Job Description:
{job_desc}
"""
    data = {
        "contents": [{"parts": [{"text": prompt}]}]
    }
    try:
        response = requests.post(url, json=data)
        response.raise_for_status()
        result = response.json()
        # Extract the improved resume text from the API response
        return result['candidates'][0]['content']['parts'][0]['text']
    except requests.exceptions.HTTPError as http_err:
        print("HTTP error occurred:", http_err)
        try:
            print("Response content:", response.content.decode())
        except Exception:
            print("Could not decode response content.")
        raise
    except Exception as err:
        print("Other error occurred:", err)
        try:
            print("Response content:", response.content.decode())
        except Exception:
            print("Could not decode response content.")
        raise

def update_docx_with_text(doc, new_text):
    """
    Update the Document object with new text, preserving formatting by editing text within runs.
    Each line in new_text replaces the corresponding paragraph's text.
    """
    lines = new_text.split('\n')
    for i, para in enumerate(doc.paragraphs):
        if i < len(lines):
            new_line = lines[i]
            # If paragraph has runs, update text in runs
            if para.runs:
                # Join all run texts to get original text length
                total_len = sum(len(run.text) for run in para.runs)
                # Split new_line into chunks matching run lengths
                start = 0
                for run in para.runs:
                    run_len = len(run.text)
                    run.text = new_line[start:start+run_len]
                    start += run_len
                # If new_line is longer, append to last run
                if start < len(new_line) and para.runs:
                    para.runs[-1].text += new_line[start:]
            else:
                para.text = new_line
        else:
            # Leave extra paragraphs unchanged to preserve formatting
            pass
    # If more lines, add new paragraphs
    for line in lines[len(doc.paragraphs):]:
        doc.add_paragraph(line)
    return doc

def main():

    print("--- Resume Auto-Updater with Gemini ---")
    # Usage: python main.py resume.docx jobdesc.txt [API_KEY]
    if len(sys.argv) < 3:
        print("Usage: python main.py <resume.docx> <jobdesc.txt|jobdesc.docx> [API_KEY]")
        sys.exit(1)
    resume_path = sys.argv[1]
    job_desc_path = sys.argv[2]
    if len(sys.argv) >= 4:
        api_key = sys.argv[3]
    else:
        api_key = get_gemini_api_key()

    resume_text, doc = read_docx_text(resume_path)
    if job_desc_path.lower().endswith('.docx'):
        job_desc, _ = read_docx_text(job_desc_path)
    else:
        job_desc = read_txt_file(job_desc_path)

    # Parse docx into sections
    section_map = parse_docx_sections(doc)
    # Set word limits for each section (adjust as needed)
    word_limits = {'intro': 70, 'skills': 60, 'work_experience': 250}
    # Extract editable sections
    editable_sections = {}
    for sec in ['intro', 'skills', 'work_experience']:
        indices = section_map[sec]
        if indices:
            editable_sections[sec] = extract_section_text(doc, indices, word_limit=word_limits[sec])
        else:
            editable_sections[sec] = ''

    # Prepare prompt for Gemini: only editable sections
    resume_for_gemini = (
        f"INTRO PARAGRAPH:\n{editable_sections['intro']}\n\n"
        f"SKILLS SECTION:\n{editable_sections['skills']}\n\n"
        f"WORK EXPERIENCE SECTION:\n{editable_sections['work_experience']}\n"
    )

    print("Contacting Gemini API to update your resume...")
    new_resume_text = call_gemini(api_key, resume_for_gemini, job_desc)


    # Split Gemini output back into sections (assume same order)
    import re
    intro_new, skills_new, work_new = '', '', ''
    m = re.split(r'INTRO PARAGRAPH:|SKILLS SECTION:|WORK EXPERIENCE SECTION:', new_resume_text, flags=re.I)
    if len(m) >= 4:
        intro_new = m[1].strip()
        skills_new = m[2].strip()
        work_new = m[3].strip()
    else:
        work_new = new_resume_text.strip()

    # Update only allowed sections in docx if indices are present
    if section_map['intro']:
        update_section_text(doc, section_map['intro'], intro_new)
    if section_map['skills']:
        skills_bullets = process_core_competencies(skills_new)
        update_core_competencies(doc, section_map['skills'], skills_bullets)
    if section_map['work_experience']:
        work_processed = process_work_experience(work_new)
        update_work_experience(doc, section_map['work_experience'], work_processed)

    updated_docx_path = os.path.splitext(resume_path)[0] + "_updated.docx"
    doc.save(updated_docx_path)
    print(f"Updated .docx saved to: {updated_docx_path}")

    print("Converting to PDF...")
    updated_pdf_path = os.path.splitext(resume_path)[0] + ".pdf"
    convert(updated_docx_path, updated_pdf_path)
    print(f"PDF saved to: {updated_pdf_path}")

if __name__ == "__main__":
    main()
